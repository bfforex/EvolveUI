import os
import logging
from typing import Dict, Any, Optional
import uuid
from pathlib import Path

# Try to import file processing libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PyPDF2 not available. PDF processing disabled.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX processing disabled.")

try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
    logging.warning("python-magic not available. File type detection limited.")

# Set MAGIC_AVAILABLE to False for now since we can't install it
MAGIC_AVAILABLE = False

logger = logging.getLogger(__name__)

class FileProcessor:
    def __init__(self, upload_dir: str = "/tmp/uploads"):
        """Initialize file processor with upload directory"""
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
        
        # Supported file types
        self.supported_types = {
            '.txt': 'text/plain',
            '.md': 'text/markdown',
            '.py': 'text/python',
            '.js': 'text/javascript',
            '.json': 'application/json',
            '.csv': 'text/csv',
        }
        
        if PDF_AVAILABLE:
            self.supported_types['.pdf'] = 'application/pdf'
        
        if DOCX_AVAILABLE:
            self.supported_types['.docx'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    def get_supported_types(self) -> Dict[str, str]:
        """Get dictionary of supported file extensions and their MIME types"""
        return self.supported_types.copy()

    def detect_file_type(self, file_path: Path) -> str:
        """Detect file type using magic or extension"""
        if MAGIC_AVAILABLE:
            try:
                mime_type = magic.from_file(str(file_path), mime=True)
                return mime_type
            except Exception as e:
                logger.warning(f"Magic detection failed: {e}")
        
        # Fallback to extension
        extension = file_path.suffix.lower()
        return self.supported_types.get(extension, 'application/octet-stream')

    def save_uploaded_file(self, file_content: bytes, filename: str) -> Path:
        """Save uploaded file content to disk"""
        # Generate unique filename to avoid conflicts
        file_id = str(uuid.uuid4())
        extension = Path(filename).suffix
        safe_filename = f"{file_id}{extension}"
        
        file_path = self.upload_dir / safe_filename
        
        with open(file_path, 'wb') as f:
            f.write(file_content)
        
        return file_path

    def extract_text_from_file(self, file_path: Path, filename: str = None) -> Dict[str, Any]:
        """Extract text content from various file types"""
        try:
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            extension = file_path.suffix.lower()
            original_name = filename or file_path.name
            
            result = {
                "filename": original_name,
                "file_path": str(file_path),
                "file_type": extension,
                "mime_type": self.detect_file_type(file_path),
                "content": "",
                "metadata": {},
                "success": True,
                "error": None
            }
            
            if extension == '.pdf':
                result.update(self._extract_pdf_text(file_path))
            elif extension == '.docx':
                result.update(self._extract_docx_text(file_path))
            elif extension in ['.txt', '.md', '.py', '.js', '.json', '.csv']:
                result.update(self._extract_plain_text(file_path))
            else:
                result.update({
                    "success": False,
                    "error": f"Unsupported file type: {extension}"
                })
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            return {
                "filename": filename or str(file_path),
                "file_path": str(file_path),
                "file_type": file_path.suffix.lower() if file_path else "unknown",
                "content": "",
                "metadata": {},
                "success": False,
                "error": str(e)
            }

    def _extract_pdf_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF file"""
        if not PDF_AVAILABLE:
            return {
                "success": False,
                "error": "PDF processing not available. Install PyPDF2."
            }
        
        try:
            text_content = []
            metadata = {"pages": 0}
            
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                metadata["pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1}: {e}")
                
                # Get PDF metadata
                if pdf_reader.metadata:
                    metadata.update({
                        "title": pdf_reader.metadata.get('/Title', ''),
                        "author": pdf_reader.metadata.get('/Author', ''),
                        "subject": pdf_reader.metadata.get('/Subject', ''),
                        "creator": pdf_reader.metadata.get('/Creator', ''),
                    })
            
            return {
                "content": "\n\n".join(text_content),
                "metadata": metadata
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"PDF processing error: {str(e)}"
            }

    def _extract_docx_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            return {
                "success": False,
                "error": "DOCX processing not available. Install python-docx."
            }
        
        try:
            doc = Document(file_path)
            
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract table content
            tables_content = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):  # Only add non-empty rows
                        table_data.append(" | ".join(row_data))
                if table_data:
                    tables_content.append("\n".join(table_data))
            
            content_parts = []
            if paragraphs:
                content_parts.append("\n\n".join(paragraphs))
            if tables_content:
                content_parts.append("\n\n--- Tables ---\n\n" + "\n\n".join(tables_content))
            
            metadata = {
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables)
            }
            
            # Try to get document properties
            try:
                core_props = doc.core_properties
                metadata.update({
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else "",
                })
            except Exception:
                pass
            
            return {
                "content": "\n\n".join(content_parts),
                "metadata": metadata
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"DOCX processing error: {str(e)}"
            }

    def _extract_plain_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from plain text files"""
        try:
            encodings = ['utf-8', 'utf-16', 'iso-8859-1', 'cp1252']
            content = ""
            used_encoding = "utf-8"
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                        used_encoding = encoding
                        break
                except UnicodeDecodeError:
                    continue
            
            if not content:
                # Last resort: read as binary and decode with errors='replace'
                with open(file_path, 'rb') as f:
                    raw_content = f.read()
                    content = raw_content.decode('utf-8', errors='replace')
                    used_encoding = "utf-8 (with errors replaced)"
            
            metadata = {
                "encoding": used_encoding,
                "size_bytes": file_path.stat().st_size,
                "line_count": content.count('\n') + 1 if content else 0
            }
            
            return {
                "content": content,
                "metadata": metadata
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Text extraction error: {str(e)}"
            }

    def cleanup_file(self, file_path: Path) -> bool:
        """Remove uploaded file from disk"""
        try:
            if file_path.exists():
                file_path.unlink()
                return True
            return False
        except Exception as e:
            logger.error(f"Error cleaning up file {file_path}: {e}")
            return False